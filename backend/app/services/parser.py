"""
Document parsing service for PDF, DOCX, and TXT files.
"""

from pathlib import Path
from typing import List

import fitz  # PyMuPDF
import pdfplumber
from docx import Document as DocxDocument


class DocumentParser:
    """Parses documents and extracts text."""

    @staticmethod
    def parse(file_path: str, doc_type: str) -> List[str]:
        """
        Parse a document and return list of text sections.

        Args:
            file_path: Path to the file
            doc_type: Type of document (pdf, docx, txt)

        Returns:
            List of text strings (one per page/section)

        Raises:
            ValueError: If document type is not supported
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if doc_type.lower() in ["pdf", "brd", "api_spec", "qa_policy"]:
            return DocumentParser._parse_pdf(file_path)
        elif doc_type.lower() in ["docx", "doc", "user_story"]:
            return DocumentParser._parse_docx(file_path)
        elif doc_type.lower() in ["txt", "nfr"]:
            return DocumentParser._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")

    @staticmethod
    def _parse_pdf(file_path: str) -> List[str]:
        """
        Parse PDF using PyMuPDF (fitz).
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            List of text pages
        """
        pages = []
        try:
            pdf_doc = fitz.open(file_path)
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                text = page.get_text()
                if text.strip():  # Only add non-empty pages
                    pages.append(text)
            pdf_doc.close()
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")

        if not pages:
            raise ValueError("PDF contains no readable text")

        return pages

    @staticmethod
    def _parse_docx(file_path: str) -> List[str]:
        """
        Parse DOCX using python-docx.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of paragraphs and table contents
        """
        sections = []
        try:
            doc = DocxDocument(file_path)

            for para in doc.paragraphs:
                if para.text.strip():
                    sections.append(para.text)

            # Extract table data
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells
                    )
                    if row_text.strip():
                        sections.append(row_text)

        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")

        if not sections:
            raise ValueError("DOCX contains no readable content")

        return sections

    @staticmethod
    def _parse_txt(file_path: str) -> List[str]:
        """
        Parse TXT file.
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            List with entire file content (and optionally split by sections)
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            if not content.strip():
                raise ValueError("TXT file is empty")

            # Split by double newlines to preserve some structure
            sections = [s.strip() for s in content.split("\n\n") if s.strip()]
            return sections if sections else [content]

        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, "r", encoding="latin-1") as f:
                    content = f.read()
                sections = [s.strip() for s in content.split("\n\n") if s.strip()]
                return sections if sections else [content]
            except Exception as e:
                raise ValueError(f"Error parsing TXT: {str(e)}")
        except Exception as e:
            raise ValueError(f"Error parsing TXT: {str(e)}")
